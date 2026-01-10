import streamlit as st
import pandas as pd
import numpy as np
from pathlib import Path
import base64
from io import BytesIO
from datetime import date
from typing import Optional, Tuple, List, Dict
import textwrap
import re
import json
import hashlib

# DOCX (python-docx)
try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
except Exception:
    Document = None


# ------------------------------------------------------------
# Config
# ------------------------------------------------------------
st.set_page_config(page_title="Booking Platform", layout="wide")

BASE_DIR = Path(__file__).parent
ASSETS_DIR = BASE_DIR / "assets"
DATA_DIR = BASE_DIR / "data"
DOOH_MASTER_PATH = DATA_DIR / "dooh_master.csv"


# ------------------------------------------------------------
# Sales contacts (dropdown)
# ------------------------------------------------------------
SALES_CONTACTS = [
    ("Darren van der Schyff", "darren@vicinity-media.com"),
    ("Dwane McCarthy", "Dwane@vicinity-media.com"),
    ("Frankie Matianyi", "frankie@vicinity-media.com"),
    ("Greg Sinnett", "greg@vicinity-media.com"),
    ("Nabeel Haroon", "nabeel@vicinity-media.com"),
    ("Ollie Westphal", "ollie@vicinity-media.com"),
]
SALES_EMAIL_BY_NAME = {n: e for n, e in SALES_CONTACTS}
SALES_NAMES = [n for n, _ in SALES_CONTACTS]


# ------------------------------------------------------------
# Helpers: safe CSV read from UploadedFile (fixes “works local, not on Cloud” seek quirks)
# ------------------------------------------------------------
def read_csv_uploaded(file, **kwargs) -> pd.DataFrame:
    if file is None:
        return pd.DataFrame()
    try:
        b = file.getvalue()
        return pd.read_csv(BytesIO(b), **kwargs)
    except Exception:
        # fallback (if getvalue not available for some reason)
        try:
            file.seek(0)
            return pd.read_csv(file, **kwargs)
        except Exception:
            return pd.DataFrame()


# ------------------------------------------------------------
# Helpers: assets + styling
# ------------------------------------------------------------
def find_asset(stem: str) -> Optional[Path]:
    """Find an asset by stem in /assets (with or without extension)."""
    if not ASSETS_DIR.exists():
        return None

    exact = ASSETS_DIR / stem
    if exact.exists() and exact.is_file():
        return exact

    matches = list(ASSETS_DIR.glob(f"{stem}.*"))
    return matches[0] if matches else None


def sniff_mime_type(path: Path) -> str:
    header = path.read_bytes()[:16]
    if header.startswith(b"\x89PNG\r\n\x1a\n"):
        return "image/png"
    if header.startswith(b"\xff\xd8\xff"):
        return "image/jpeg"
    if header.startswith(b"RIFF") and b"WEBP" in header:
        return "image/webp"
    return "image/png"


def file_to_base64(path: Path) -> str:
    return base64.b64encode(path.read_bytes()).decode("utf-8")


def inject_background(background_path: Path, white_overlay_opacity: float = 0.45):
    """Full-page background image + white overlay."""
    b64 = file_to_base64(background_path)
    mime = sniff_mime_type(background_path)

    st.markdown(
        f"""
        <style>
          .stApp {{
            background-image: url("data:{mime};base64,{b64}");
            background-size: cover;
            background-position: center top;
            background-repeat: no-repeat;
            background-attachment: fixed;
          }}

          /* white overlay */
          .stApp::before {{
            content: "";
            position: fixed;
            inset: 0;
            background: rgba(255,255,255,{white_overlay_opacity});
            pointer-events: none;
            z-index: 0;
          }}

          /* keep content above overlay */
          section[data-testid="stSidebar"], .main, header {{
            position: relative;
            z-index: 2;
          }}
        </style>
        """,
        unsafe_allow_html=True,
    )


def inject_app_css():
    """
    Fix logo/header clipping + KPI card styling
    """
    st.markdown(
        """
        <style>
          header[data-testid="stHeader"] { height: 0px !important; }
          header[data-testid="stHeader"] * { display: none !important; }
          div[data-testid="stToolbar"] { display: none !important; }
          #MainMenu { visibility: hidden !important; }
          footer { visibility: hidden !important; }

          .stApp, .main, section.main, section.main > div, div[data-testid="stAppViewContainer"] {
            overflow: visible !important;
          }
          div[data-testid="stMarkdown"], div[data-testid="stMarkdownContainer"] {
            overflow: visible !important;
          }

          input[type="radio"], input[type="checkbox"] {
            accent-color: #0B2A4A !important;
          }

          @keyframes fadeInUp {
            from { opacity: 0; transform: translateY(10px); }
            to   { opacity: 1; transform: translateY(0); }
          }
          .fade-in { animation: fadeInUp 520ms ease-out both; }

          :root { --sticky-h: 240px; }

          .sticky-wrap {
            position: fixed;
            top: 0; left: 0; right: 0;
            z-index: 999999;
            background: rgba(255,255,255,0.97);
            backdrop-filter: blur(6px);
            overflow: visible;
            box-sizing: border-box;
            min-height: var(--sticky-h);
            padding: 18px 0 14px 0;
          }

          .logo-bar { width: 100%; margin: 0; padding: 0; overflow: visible; box-sizing: border-box; }
          .logo-inner {
            display: flex;
            justify-content: center;
            align-items: center;
            overflow: visible;
            padding: 6px 0 0 0;
            box-sizing: border-box;
          }

          .logo-inner img {
            display: block;
            width: min(460px, 78vw);
            height: auto;
            max-height: 120px;
            object-fit: contain;
            margin: 0;
            padding: 0;
          }

          .logo-divider {
            height: 1px;
            width: 100%;
            background: rgba(15, 23, 42, 0.08);
            margin-top: 12px;
          }

          .title-wrap { text-align: center; padding: 10px 0 0 0; }
          .app-title { color: #0B2A4A; font-size: 24px; font-weight: 800; margin: 0; line-height: 1.2; }

          /* Map. Budget. Book. made larger */
          .app-caption { color: #475569; font-size: 16px; font-weight: 700; margin-top: 6px; line-height: 1.2; }

          .block-container { padding-top: calc(var(--sticky-h) + 16px) !important; }

          div[data-testid="stDataFrame"] { border-radius: 10px; }

          .kpi-grid {
            display: grid;
            grid-template-columns: repeat(4, minmax(0, 1fr));
            gap: 14px;
            width: 100%;
            margin-top: 10px;
          }

          @media (max-width: 1200px) { .kpi-grid { grid-template-columns: repeat(3, minmax(0, 1fr)); } :root { --sticky-h: 250px; } }
          @media (max-width: 900px)  { .kpi-grid { grid-template-columns: repeat(2, minmax(0, 1fr)); } :root { --sticky-h: 260px; } }
          @media (max-width: 600px)  {
            .kpi-grid { grid-template-columns: 1fr; }
            :root { --sticky-h: 280px; }
            .app-title { font-size: 20px; }
            .app-caption { font-size: 15px; }
            .logo-inner img { width: min(380px, 85vw); max-height: 110px; }
          }

          .kpi-card {
            background: rgba(255,255,255,0.92);
            border: 1px solid rgba(15, 23, 42, 0.10);
            border-radius: 14px;
            padding: 14px 14px 12px 14px;
            box-shadow: 0 2px 10px rgba(15, 23, 42, 0.06);
          }
          .kpi-label {
            color: #64748b;
            font-size: 12px;
            font-weight: 600;
            letter-spacing: 0.2px;
            margin-bottom: 6px;
          }
          .kpi-value {
            color: #0B2A4A;
            font-size: 22px;
            font-weight: 800;
            line-height: 1.1;
            word-break: break-word;
          }

          .rec-box {
            background: rgba(255,255,255,0.92);
            border: 1px solid rgba(15, 23, 42, 0.10);
            border-radius: 14px;
            padding: 14px 14px 12px 14px;
            box-shadow: 0 2px 10px rgba(15, 23, 42, 0.06);
            margin-top: 12px;
          }
          .rec-title { font-weight: 800; color: #0B2A4A; margin-bottom: 6px; }
          .rec-body  { color: #334155; font-size: 13px; line-height: 1.45; }
          .rec-bullets { margin-top: 8px; color: #334155; font-size: 13px; line-height: 1.45; }
          .rec-bullets li { margin-bottom: 6px; }

          /* NEW: bold callout for IO instruction on Tab 3 (ensures it renders clearly on Streamlit Cloud) */
          .io-callout {
            background: rgba(11, 42, 74, 0.06);
            border: 1px solid rgba(11, 42, 74, 0.14);
            border-radius: 14px;
            padding: 14px 16px;
            margin: 10px 0 14px 0;
            color: #0B2A4A;
            font-size: 14px;
            line-height: 1.5;
            font-weight: 800;
          }
        </style>
        """,
        unsafe_allow_html=True,
    )


def render_header(logo_path: Optional[Path]):
    logo_html = ""
    if logo_path and logo_path.exists():
        b64 = file_to_base64(logo_path)
        mime = sniff_mime_type(logo_path)
        logo_html = f'<img src="data:{mime};base64,{b64}" alt="Vicinity Logo" />'

    st.markdown(
        f"""
        <div class="sticky-wrap fade-in">
          <div class="logo-bar">
            <div class="logo-inner">{logo_html}</div>
            <div class="logo-divider"></div>
          </div>

          <div class="title-wrap">
            <div class="app-title">The Proximity-First Media Planning Engine</div>
            <div class="app-caption">Map. Budget. Book.</div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def kpi_grid(items: List[Dict[str, str]]):
    cards_html = []
    for it in items:
        label = it.get("label", "")
        value = it.get("value", "")
        card = textwrap.dedent(
            f"""
            <div class="kpi-card">
              <div class="kpi-label">{label}</div>
              <div class="kpi-value">{value}</div>
            </div>
            """
        ).strip()
        cards_html.append(card)

    grid_html = "<div class='kpi-grid'>" + "".join(cards_html) + "</div>"
    st.markdown(grid_html, unsafe_allow_html=True)


def render_recommendation_box(title: str, body: str, bullets: List[str]):
    bullets_html = "".join([f"<li>{b}</li>" for b in bullets if b])
    st.markdown(
        f"""
        <div class="rec-box">
          <div class="rec-title">{title}</div>
          <div class="rec-body">{body}</div>
          <ul class="rec-bullets">{bullets_html}</ul>
        </div>
        """,
        unsafe_allow_html=True,
    )


# ------------------------------------------------------------
# Helpers: column detection
# ------------------------------------------------------------
def pick_lat_lon(df: pd.DataFrame) -> Tuple[Optional[str], Optional[str]]:
    cols = [c.strip() for c in df.columns]
    lower = {c: c.lower() for c in cols}

    lat_candidates = [c for c in cols if "lat" in lower[c]]
    lon_candidates = [c for c in cols if ("lon" in lower[c]) or ("lng" in lower[c]) or ("long" in lower[c])]

    if not lat_candidates or not lon_candidates:
        return None, None
    return lat_candidates[0], lon_candidates[0]


def pick_name_col(df: pd.DataFrame) -> Optional[str]:
    priority = [
        "branch_name", "branch", "store_name", "store", "name", "location",
        "site_name", "site", "title", "outlet", "outlet_name", "id"
    ]
    low_map = {c.lower(): c for c in df.columns}
    for p in priority:
        if p in low_map:
            return low_map[p]

    for c in df.columns:
        cl = c.lower()
        if ("lat" in cl) or ("lon" in cl) or ("lng" in cl) or ("long" in cl):
            continue
        return c
    return None


def pick_first_existing(df: pd.DataFrame, candidates: List[str]) -> Optional[str]:
    low_map = {c.lower().strip(): c for c in df.columns}
    for cand in candidates:
        key = cand.lower().strip()
        if key in low_map:
            return low_map[key]
    return None


def pick_province_col(df: pd.DataFrame) -> Optional[str]:
    return pick_first_existing(df, ["Province", "province", "PROVINCE"])


# ------------------------------------------------------------
# DOOH master loading + Province normalization
# ------------------------------------------------------------
def normalize_province_name(x: str) -> str:
    if x is None or (isinstance(x, float) and np.isnan(x)):
        return ""
    s = str(x).strip()
    if not s:
        return ""

    s = re.sub(r"\s+", " ", s)
    s = s.replace("–", "-").replace("—", "-")
    s_low = s.lower()

    mapping = {
        "kwazulu natal": "KwaZulu-Natal",
        "kwazulu-natal": "KwaZulu-Natal",
        "kzn": "KwaZulu-Natal",
        "north west": "North West",
        "north-west": "North West",
        "northwest": "North West",
        "eastern cape": "Eastern Cape",
        "western cape": "Western Cape",
        "northern cape": "Northern Cape",
        "free state": "Free State",
        "gauteng": "Gauteng",
        "limpopo": "Limpopo",
        "mpumalanga": "Mpumalanga",
    }
    if s_low in mapping:
        return mapping[s_low]

    parts = []
    for chunk in s.split(" "):
        if "-" in chunk:
            hy = "-".join([p[:1].upper() + p[1:].lower() if p else "" for p in chunk.split("-")])
            parts.append(hy)
        else:
            parts.append(chunk[:1].upper() + chunk[1:].lower())
    out = " ".join(parts)

    if out.lower() in ["kwazulu-natal", "kwazulu natal"]:
        out = "KwaZulu-Natal"
    if out.lower() in ["north west", "north-west", "northwest"]:
        out = "North West"
    return out


def load_dooh_master() -> Optional[pd.DataFrame]:
    """Always read master fresh from disk (no caching)."""
    if not DOOH_MASTER_PATH.exists():
        return None
    try:
        df = pd.read_csv(DOOOH_MASTER_PATH)  # intentionally incorrect? (NO) -> fix below
        return df
    except Exception:
        return None


# Fix: keep original function name/behavior (no unintended changes)
def load_dooh_master() -> Optional[pd.DataFrame]:
    """Always read master fresh from disk (no caching)."""
    if not DOOH_MASTER_PATH.exists():
        return None
    try:
        df = pd.read_csv(DOOH_MASTER_PATH)
        prov_col = pick_province_col(df)
        if prov_col:
            df[prov_col] = df[prov_col].apply(normalize_province_name)
        return df
    except Exception:
        return None


# ------------------------------------------------------------
# Distance
# ------------------------------------------------------------
def haversine_vec(lat1, lon1, lat2_arr, lon2_arr):
    R = 6371.0
    lat1, lon1 = np.radians(lat1), np.radians(lon1)
    lat2 = np.radians(lat2_arr)
    lon2 = np.radians(lon2_arr)

    dlat = lat2 - lat1
    dlon = lon2 - lon1

    a = np.sin(dlat / 2) ** 2 + np.cos(lat1) * np.cos(lat2) * np.sin(dlon / 2) ** 2
    return 2 * R * np.arcsin(np.sqrt(a))


# ------------------------------------------------------------
# Formatting helpers
# ------------------------------------------------------------
def fmt_currency_rands(x: float) -> str:
    try:
        return f"R {x:,.2f}"
    except Exception:
        return "R 0.00"


def fmt_int(x: float) -> str:
    try:
        return f"{int(round(x)):,}"
    except Exception:
        return "0"


# ------------------------------------------------------------
# Templates (CSV downloads)
# ------------------------------------------------------------
def csv_bytes_from_df(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode("utf-8")


def stores_template_df() -> pd.DataFrame:
    return pd.DataFrame([{
        "Store Name": "Example Store",
        "Latitude": -26.2041,
        "Longitude": 28.0473,
        "Province": "Gauteng",
    }])


def dooh_selection_template_df() -> pd.DataFrame:
    return pd.DataFrame([{
        "Site ID/Number": "ALV 01",
        "Selected": 1,
    }])


def mobile_locations_template_df() -> pd.DataFrame:
    return pd.DataFrame([{
        "Location Name": "Example Location",
        "Latitude": -26.2041,
        "Longitude": 28.0473,
        "Province": "Gauteng",
    }])


# ------------------------------------------------------------
# Selection / counting logic
# ------------------------------------------------------------
def load_selected_sites_from_upload(file) -> pd.DataFrame:
    df = read_csv_uploaded(file)
    site_col = pick_first_existing(df, ["Site ID/Number", "Site ID", "SiteID", "SiteID/Number", "Site_Number", "Site Number"])
    sel_col = pick_first_existing(df, ["Selected", "selected", "SELECTED"])

    if not site_col or not sel_col:
        raise ValueError("Selection file must contain 'Site ID/Number' and 'Selected' columns.")

    df[sel_col] = df[sel_col].astype(str).str.strip().str.lower()
    df["_selected_flag"] = df[sel_col].isin(["1", "true", "yes", "y"])
    selected = df[df["_selected_flag"]].copy()
    selected = selected[[site_col]].rename(columns={site_col: "Site ID/Number"})
    selected["Site ID/Number"] = selected["Site ID/Number"].astype(str).str.strip()
    selected = selected.dropna().drop_duplicates()
    return selected


def count_mobile_locations_any_format(file) -> int:
    df = read_csv_uploaded(file)

    sel_col = pick_first_existing(df, ["Selected", "selected", "SELECTED"])
    site_col = pick_first_existing(df, ["Site ID/Number", "Site ID", "SiteID", "Site Number"])

    if sel_col and site_col:
        tmp = df.copy()
        tmp[sel_col] = tmp[sel_col].astype(str).str.strip().str.lower()
        tmp["_selected_flag"] = tmp[sel_col].isin(["1", "true", "yes", "y"])
        return int(tmp["_selected_flag"].sum())

    lat, lon = pick_lat_lon(df)
    if not lat or not lon:
        return int(len(df))

    df[lat] = pd.to_numeric(df[lat], errors="coerce")
    df[lon] = pd.to_numeric(df[lon], errors="coerce")
    df = df.dropna(subset=[lat, lon])
    return int(len(df))


def dooh_cpm_from_count(n_locations: int) -> int:
    return 240 if n_locations >= 50 else 260


def mobile_cpm_from_count(n_locations: int) -> int:
    return 140 if n_locations >= 50 else 160


def impressions_from_budget_and_cpm(budget: float, cpm: float) -> float:
    if cpm <= 0:
        return 0
    return (budget / cpm) * 1000.0


# ------------------------------------------------------------
# Savings message helper
# ------------------------------------------------------------
def savings_message(product: str, n_locs: int, budget: float) -> str:
    product = (product or "").strip()

    if budget <= 0 or n_locs >= 50:
        return ""

    if product == "DOOH":
        high_cpm, low_cpm = 260, 240
        unit_word = "DOOH placements"
    else:
        high_cpm, low_cpm = 160, 140
        unit_word = "locations"

    current_imps = impressions_from_budget_and_cpm(budget, high_cpm)
    budget_at_low = (current_imps / 1000.0) * low_cpm

    savings_r = budget - budget_at_low
    savings_pct = (savings_r / budget) * 100 if budget > 0 else 0

    imps_at_low = impressions_from_budget_and_cpm(budget, low_cpm)
    extra_imps = imps_at_low - current_imps

    needed = max(0, 50 - int(n_locs))

    return (
        f"**Potential saving:** If you increased your {unit_word} by **{needed}** (to reach **50+**), "
        f"your CPM would drop from **R {high_cpm}** to **R {low_cpm}**. "
        f"For the *same estimated delivery*, you could save about **{fmt_currency_rands(savings_r)}** "
        f"(**{savings_pct:.1f}%**). Alternatively, keeping the same budget could add roughly **{fmt_int(extra_imps)}** impressions."
    )


# ------------------------------------------------------------
# Proximity insights (rule-based, no AI)
# ------------------------------------------------------------
def compute_insight_stats(summary_df: pd.DataFrame, pairwise_df: Optional[pd.DataFrame], mode: str, radius_km: int):
    if summary_df is None or summary_df.empty:
        return None

    cov_col = f"{'DOOH' if mode == 'Store to DOOH' else 'Store'} sites within {radius_km}km"
    if cov_col not in summary_df.columns:
        return None

    total_locations = int(len(summary_df))
    covered_mask = (pd.to_numeric(summary_df[cov_col], errors="coerce").fillna(0) > 0)
    covered_locations = int(covered_mask.sum())
    coverage_pct = int(round((covered_locations / total_locations) * 100, 0)) if total_locations else 0
    uncovered_locations = int(total_locations - covered_locations)

    nearest_col = f"Nearest {'DOOH' if mode == 'Store to DOOH' else 'Store'} distance (km)"
    nearest_vals_all = pd.to_numeric(summary_df[nearest_col], errors="coerce").dropna()
    nearest_vals_cov = pd.to_numeric(summary_df.loc[covered_mask, nearest_col], errors="coerce").dropna()

    median_nearest_cov = float(nearest_vals_cov.median()) if len(nearest_vals_cov) else None
    p90_nearest_cov = float(np.percentile(nearest_vals_cov.to_numpy(), 90)) if len(nearest_vals_cov) else None

    overall_min_dist = float(nearest_vals_all.min()) if len(nearest_vals_all) else None
    overall_min_row = None
    if overall_min_dist is not None and np.isfinite(overall_min_dist):
        try:
            overall_min_row = summary_df.loc[pd.to_numeric(summary_df[nearest_col], errors="coerce").idxmin()]
        except Exception:
            overall_min_row = None

    avg_sites = float(pd.to_numeric(summary_df[cov_col], errors="coerce").fillna(0).mean()) if total_locations else 0.0

    opp_locs = summary_df.loc[~covered_mask, "Location"].astype(str).tolist()
    opp_preview = opp_locs[:100]

    isolated_detail = []
    if mode == "Store to Store":
        near_name_col = "Nearest Store"
        if near_name_col in summary_df.columns:
            tmp_iso = summary_df.loc[~covered_mask, ["Location", nearest_col, near_name_col]].copy()
            tmp_iso[nearest_col] = pd.to_numeric(tmp_iso[nearest_col], errors="coerce")
            tmp_iso = tmp_iso.sort_values(nearest_col, ascending=True)
            for _, r in tmp_iso.head(50).iterrows():
                isolated_detail.append({
                    "Location": str(r.get("Location", "")),
                    "Nearest Store": str(r.get(near_name_col, "")),
                    "Distance (km)": float(r.get(nearest_col, np.nan)) if pd.notna(r.get(nearest_col)) else None,
                })

    top_networks = None
    if mode == "Store to DOOH" and isinstance(pairwise_df, pd.DataFrame) and not pairwise_df.empty:
        if ("Network" in pairwise_df.columns) and ("Site ID/Number" in pairwise_df.columns):
            tmp = pairwise_df.copy()
            tmp["Network"] = tmp["Network"].astype(str).str.strip()
            tmp["Site ID/Number"] = tmp["Site ID/Number"].astype(str).str.strip()
            agg = (
                tmp.dropna(subset=["Network", "Site ID/Number"])
                   .drop_duplicates(subset=["Network", "Site ID/Number"])
                   .groupby("Network", dropna=False)["Site ID/Number"]
                   .nunique()
                   .reset_index(name="Unique sites (within radius)")
                   .sort_values("Unique sites (within radius)", ascending=False)
                   .head(10)
                   .rename(columns={"Network": "DOOH Network"})
            )
            top_networks = agg

    # NEW: Province insights (only if Province exists and has values)
    province_insights = None
    if "Province" in summary_df.columns:
        prov = summary_df["Province"].astype(str).fillna("").str.strip()
        has_prov = (prov != "") & (prov.str.lower() != "nan")
        if has_prov.any():
            tmp = summary_df.copy()
            tmp["_prov"] = prov.where(has_prov, "")
            tmp = tmp[tmp["_prov"] != ""].copy()
            tmp["_covered"] = covered_mask.loc[tmp.index].astype(int)

            byp = (
                tmp.groupby("_prov", dropna=False)
                   .agg(
                        locations=("Location", "count"),
                        covered=("_covered", "sum"),
                        avg_sites=(cov_col, lambda x: float(pd.to_numeric(x, errors="coerce").fillna(0).mean()))
                   )
                   .reset_index()
                   .rename(columns={"_prov": "Province"})
            )
            byp["coverage_pct"] = np.where(
                byp["locations"] > 0,
                (byp["covered"] / byp["locations"]) * 100.0,
                0.0
            )

            # Add isolated count for Store-to-Store
            if mode == "Store to Store":
                byp["isolated"] = byp["locations"] - byp["covered"]
                byp["isolated_pct"] = np.where(
                    byp["locations"] > 0,
                    (byp["isolated"] / byp["locations"]) * 100.0,
                    0.0
                )

            province_insights = byp.sort_values(["locations", "coverage_pct"], ascending=[False, False]).reset_index(drop=True)

    return {
        "mode": mode,
        "radius_km": int(radius_km),
        "locations_analysed": total_locations,
        "coverage_pct": int(coverage_pct),
        "covered_locations": covered_locations,
        "opportunity_windows": uncovered_locations,
        "median_nearest_km": median_nearest_cov,
        "p90_nearest_km": p90_nearest_cov,
        "avg_sites_per_location": avg_sites,
        "opportunity_locations_preview": opp_preview,
        "isolated_detail": isolated_detail,
        "overall_min_row": overall_min_row.to_dict() if hasattr(overall_min_row, "to_dict") else None,
        "top_networks": top_networks,
        "province_insights": province_insights,
    }


def build_vicinity_recommendation_dooh(stats: Dict) -> Tuple[str, str, List[str]]:
    coverage_pct = int(stats["coverage_pct"])
    radius_km = int(stats["radius_km"])
    total_locations = int(stats["locations_analysed"])
    covered_locations = int(stats["covered_locations"])
    opp = int(stats["opportunity_windows"])
    median_nearest = stats.get("median_nearest_km")
    p90_nearest = stats.get("p90_nearest_km")
    avg_sites = float(stats.get("avg_sites_per_location", 0.0))

    if coverage_pct < 40:
        label = "Low coverage"
        meaning = "DOOH availability is limited at this radius, which will constrain reach and frequency if you rely on DOOH alone."
    elif coverage_pct <= 50:
        label = "Fair coverage"
        meaning = "Coverage is uneven: you have meaningful DOOH in some areas, but clear gaps that will weaken consistent delivery."
    elif coverage_pct <= 59:
        label = "Moderate coverage"
        meaning = "Coverage is workable, but you’ll still have pockets where DOOH won’t consistently support delivery at this radius."
    elif coverage_pct <= 75:
        label = "High coverage"
        meaning = "DOOH availability is strong across most locations, which supports efficient reach planning at this radius."
    else:
        label = "Very high coverage"
        meaning = "Coverage is excellent and highly consistent, giving you strong DOOH density around most locations."

    title = "Vicinity recommendation"
    body = (
        f"At {radius_km}km, {coverage_pct}% of your locations have nearby DOOH coverage "
        f"({covered_locations} out of {total_locations}). {label}. {meaning}"
    )

    bullets = []
    if opp > 0:
        bullets.append(
            f"Gap list: {opp} locations have zero nearby DOOH sites at {radius_km}km. Treat these as priority areas for gap-cover channels (Mobile/CTV) or radius expansion."
        )
    else:
        bullets.append("Gaps: No zero-coverage locations were identified at this radius.")

    if median_nearest is not None:
        if p90_nearest is not None:
            bullets.append(
                f"Proximity quality (covered locations only): typical nearest distance is {median_nearest:.2f}km (90% within {p90_nearest:.2f}km)."
            )
        else:
            bullets.append(f"Proximity quality (covered locations only): typical nearest distance is {median_nearest:.2f}km.")

    if coverage_pct < 60:
        bullets.append("Radius planning: increase the radius to improve coverage (e.g., +3–8km) and re-check where coverage becomes stable without diluting relevance too far.")
    else:
        bullets.append("Radius planning: keep the radius as-is for relevance, and focus on budget efficiency in the highest-coverage areas.")

    if avg_sites < 0.5:
        bullets.append(f"Inventory density: DOOH density is light (avg {avg_sites:.2f} sites per location within radius). Use Mobile/CTV to protect frequency in areas where DOOH is sparse.")
    else:
        bullets.append(f"Inventory density: DOOH density is healthy (avg {avg_sites:.2f} sites per location), supporting consistent exposure in covered areas.")

    top_nets = stats.get("top_networks")
    if isinstance(top_nets, pd.DataFrame) and not top_nets.empty:
        top_net = str(top_nets.iloc[0]["DOOH Network"])
        bullets.append(f"Where coverage is strongest: nearby inventory is most concentrated in {top_net}.")

    # NEW: Province insight (only if provinces were included)
    prov_df = stats.get("province_insights")
    if isinstance(prov_df, pd.DataFrame) and not prov_df.empty:
        tmp = prov_df.copy()
        tmp["coverage_pct"] = pd.to_numeric(tmp["coverage_pct"], errors="coerce").fillna(0.0)
        tmp["locations"] = pd.to_numeric(tmp["locations"], errors="coerce").fillna(0).astype(int)

        # pick strongest + weakest among provinces with meaningful volume
        meaningful = tmp[tmp["locations"] >= 2].copy()
        if meaningful.empty:
            meaningful = tmp.copy()

        best = meaningful.sort_values(["coverage_pct", "locations"], ascending=[False, False]).head(1)
        worst = meaningful.sort_values(["coverage_pct", "locations"], ascending=[True, False]).head(1)

        best_txt = ""
        worst_txt = ""
        if not best.empty:
            r = best.iloc[0]
            best_txt = f"{r['Province']} ({r['coverage_pct']:.0f}% coverage across {int(r['locations'])} locations)"
        if not worst.empty:
            r = worst.iloc[0]
            worst_txt = f"{r['Province']} ({r['coverage_pct']:.0f}% coverage across {int(r['locations'])} locations)"

        if best_txt and worst_txt and best_txt != worst_txt:
            bullets.append(f"Province view (where provided): strongest coverage in {best_txt}; weakest coverage in {worst_txt}.")
        elif best_txt:
            bullets.append(f"Province view (where provided): strongest coverage in {best_txt}.")

    return title, body, bullets


def build_vicinity_recommendation_store_to_store(stats: Dict) -> Tuple[str, str, List[str]]:
    radius_km = int(stats["radius_km"])
    total_locations = int(stats["locations_analysed"])
    isolated = int(stats["opportunity_windows"])
    coverage_pct = int(stats["coverage_pct"])

    overall_min_row = stats.get("overall_min_row") or {}
    overall_a = str(overall_min_row.get("Location", "")).strip()
    overall_b = str(overall_min_row.get("Nearest Store", "")).strip()
    overall_d = overall_min_row.get("Nearest Store distance (km)", None)

    title = "Vicinity recommendation"

    if overall_a and overall_b and isinstance(overall_d, (int, float, np.floating)) and np.isfinite(overall_d):
        closest_sentence = f"The closest store pairing is {overall_a} → {overall_b} at {float(overall_d):.2f}km."
    else:
        closest_sentence = "Closest-store pairing could not be reliably detected (check Store name/lat/lon columns)."

    if isolated == 0:
        isolation_sentence = f"At {radius_km}km, every store has at least one other store within the radius (no isolated stores)."
    else:
        isolation_sentence = (
            f"At {radius_km}km, {isolated} stores have no other store within the radius "
            f"({coverage_pct}% have at least one nearby store)."
        )

    body = f"{closest_sentence} {isolation_sentence}"

    bullets = []
    if isolated > 0:
        bullets.append(
            f"Isolation risk: those {isolated} stores are effectively operating in low-density catchments at {radius_km}km (nearest store exists, but sits outside the radius)."
        )
        iso_detail = stats.get("isolated_detail") or []
        if iso_detail:
            examples = iso_detail[:5]
            ex_txt = "; ".join(
                [
                    f"{x['Location']} → {x['Nearest Store']} ({x['Distance (km)']:.2f}km)"
                    for x in examples
                    if x.get("Distance (km)") is not None
                ]
            )
            if ex_txt:
                bullets.append(f"Examples (isolated stores): {ex_txt}.")
        bullets.append("Action idea: consider expanding radius (e.g., +5–15km) to see where these isolated stores begin to connect into clusters, or treat them as standalone coverage areas.")
    else:
        bullets.append("Network shape: store coverage is clustered enough that all stores connect to at least one other store within the selected radius.")
        bullets.append("Action idea: you can keep the radius tight for relevance and focus on identifying the densest clusters (high competition / overlap) vs quieter areas.")

    median_nearest = stats.get("median_nearest_km")
    p90_nearest = stats.get("p90_nearest_km")
    if median_nearest is not None:
        if p90_nearest is not None:
            bullets.append(f"Nearest-store distance (for stores with at least one nearby store): typical nearest is {median_nearest:.2f}km (90% within {p90_nearest:.2f}km).")
        else:
            bullets.append(f"Nearest-store distance (for stores with at least one nearby store): typical nearest is {median_nearest:.2f}km.")

    # NEW: Province insight (only if provinces were included)
    prov_df = stats.get("province_insights")
    if isinstance(prov_df, pd.DataFrame) and not prov_df.empty and "isolated" in prov_df.columns:
        tmp = prov_df.copy()
        tmp["isolated"] = pd.to_numeric(tmp["isolated"], errors="coerce").fillna(0).astype(int)
        tmp["locations"] = pd.to_numeric(tmp["locations"], errors="coerce").fillna(0).astype(int)

        top_iso = tmp.sort_values(["isolated", "locations"], ascending=[False, False]).head(1)
        if not top_iso.empty and int(top_iso.iloc[0]["isolated"]) > 0:
            r = top_iso.iloc[0]
            bullets.append(
                f"Province view (where provided): most isolated stores are in {r['Province']} "
                f"({int(r['isolated'])} isolated out of {int(r['locations'])})."
            )

    return title, body, bullets


# ------------------------------------------------------------
# DOCX generation helpers (FONT FIX)
# ------------------------------------------------------------
def find_default_io_template() -> Optional[Path]:
    if not DATA_DIR.exists():
        return None

    candidates = sorted(DATA_DIR.glob("*Template_Insertion Order*.docx"))
    if candidates:
        return candidates[0]

    any_docx = sorted(DATA_DIR.glob("*.docx"))
    return any_docx[0] if any_docx else None


def set_run_font(run, pt_size: int = 9, font_name: Optional[str] = "Calibri"):
    """
    Force a run to a specific font/size. Also sets East Asia / HAnsi so Word doesn't flip fonts.
    """
    try:
        run.font.size = Pt(pt_size)
        if font_name:
            run.font.name = font_name
            rFonts = run._element.rPr.rFonts
            rFonts.set(qn("w:ascii"), font_name)
            rFonts.set(qn("w:hAnsi"), font_name)
            rFonts.set(qn("w:cs"), font_name)
            rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass


def set_paragraph_font(paragraph, pt_size: int = 9, font_name: str = "Calibri"):
    """
    Ensure paragraph style isn't forcing a different size.
    - sets paragraph style font (if available)
    - sets all runs
    """
    try:
        if paragraph.style is not None and hasattr(paragraph.style, "font"):
            paragraph.style.font.name = font_name
            paragraph.style.font.size = Pt(pt_size)
    except Exception:
        pass

    try:
        for r in paragraph.runs:
            set_run_font(r, pt_size, font_name)
    except Exception:
        pass


def set_document_styles(doc: "Document", pt_size: int = 9, font_name: str = "Calibri"):
    """
    This is the main consistency fix:
    - Set Normal style default font + size
    - Also align common styles (Headings, No Spacing, Table styles) so template styles don't override.
    """
    if not hasattr(doc, "styles"):
        return

    try:
        if "Normal" in doc.styles:
            s = doc.styles["Normal"]
            if hasattr(s, "font"):
                s.font.name = font_name
                s.font.size = Pt(pt_size)
    except Exception:
        pass

    common_style_names = [
        "No Spacing",
        "Heading 1", "Heading 2", "Heading 3", "Heading 4",
        "Title", "Subtitle",
        "Table Grid", "Table Normal",
        "List Paragraph",
    ]

    for name in common_style_names:
        try:
            if name in doc.styles:
                s = doc.styles[name]
                if hasattr(s, "font"):
                    s.font.name = font_name
                    s.font.size = Pt(pt_size)
        except Exception:
            continue

    # Also harden every paragraph/table run after style defaults
    try:
        for p in doc.paragraphs:
            set_paragraph_font(p, pt_size, font_name)
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        set_paragraph_font(p, pt_size, font_name)
    except Exception:
        pass


def normalize_all_doc_fonts(doc: "Document", pt_size: int = 9, font_name: str = "Calibri"):
    """
    Backwards compatible (kept), but now calls the stronger style setter first.
    """
    set_document_styles(doc, pt_size=pt_size, font_name=font_name)


def write_value_in_paragraph_if_label(paragraph, label: str, value: str, pt_size: int = 9, font_name: str = "Calibri") -> bool:
    txt = (paragraph.text or "").strip()
    if not txt.lower().startswith(label.lower()):
        return False

    paragraph.text = label + " "
    run = paragraph.add_run(value)
    set_run_font(run, pt_size, font_name)

    # Force paragraph style/runs too (some templates have label paragraph styled differently)
    set_paragraph_font(paragraph, pt_size, font_name)

    return True


def table_contains_text(table, needle: str) -> bool:
    needle = (needle or "").lower().strip()
    if not needle:
        return False
    table_text = " ".join([(c.text or "") for row in table.rows for c in row.cells]).lower()
    return needle in table_text


def fill_block_by_table_keyword(doc: "Document", keyword: str, label_map: List[Tuple[str, str]]) -> bool:
    for table in doc.tables:
        if table_contains_text(table, keyword):
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for lbl, val in label_map:
                            if val is not None and val != "":
                                if write_value_in_paragraph_if_label(p, lbl, val, pt_size=9, font_name="Calibri"):
                                    break
            return True
    return False


def fill_media_buy_total_cell(doc: "Document", media_buy_total: float, pt_size: int = 9, font_name: str = "Calibri") -> bool:
    target_label = "media buy total"
    amount_text = fmt_currency_rands(media_buy_total)

    for table in doc.tables:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell_txt = (cell.text or "").strip().lower()
                if target_label in cell_txt:
                    if i + 1 < len(row.cells):
                        row.cells[i + 1].text = amount_text
                        for p in row.cells[i + 1].paragraphs:
                            set_paragraph_font(p, pt_size, font_name)
                        return True
    return False


def fill_top_line_numbers(doc: "Document", total_budget: float, total_impressions: int, note: str = "CPM: Mixed", pt_size: int = 9, font_name: str = "Calibri"):
    for p in doc.paragraphs:
        t = (p.text or "").strip().lower()
        if "media buy total" in t and "cpm" in t and "impressions" in t:
            p.text = f"MEDIA BUY TOTAL: {fmt_currency_rands(total_budget)} | {note} | IMPRESSIONS: {fmt_int(total_impressions)}"
            set_paragraph_font(p, pt_size, font_name)
            break


def fill_media_buy_rows(doc: "Document", line_items: List[Dict], pt_size: int = 9, font_name: str = "Calibri") -> bool:
    def normalize(s: str) -> str:
        return (s or "").strip().lower()

    for table in doc.tables:
        for r_idx, row in enumerate(table.rows):
            headers = [normalize(c.text) for c in row.cells]
            if ("product" in headers) and ("rate" in headers) and ("quantity" in headers):

                def col_idx(name_exact: str, contains: Optional[str] = None):
                    for i, h in enumerate(headers):
                        if h == name_exact:
                            return i
                    if contains:
                        for i, h in enumerate(headers):
                            if contains in h:
                                return i
                    return None

                i_product = col_idx("product")
                i_desc = col_idx("description")
                i_start = col_idx("start date")
                i_end = col_idx("end date")
                i_publisher = col_idx("publisher")
                i_targeting = col_idx("targeting")
                i_rate = col_idx("rate")
                i_metric = col_idx("metric")
                i_qty = col_idx("quantity")
                i_gross = col_idx("gross rate", contains="gross")
                i_net = col_idx("net rate(excl. agency comm)", contains="net rate")

                first_data_idx = r_idx + 1
                if first_data_idx >= len(table.rows):
                    return False

                needed = len(line_items)
                existing = len(table.rows) - first_data_idx
                while existing < needed:
                    table.add_row()
                    existing += 1

                for k, item in enumerate(line_items):
                    drow = table.rows[first_data_idx + k]

                    def set_cell(ci, val):
                        if ci is None:
                            return
                        drow.cells[ci].text = val
                        for p in drow.cells[ci].paragraphs:
                            set_paragraph_font(p, pt_size, font_name)

                    set_cell(i_product, item.get("product", ""))
                    set_cell(i_desc, item.get("description", ""))
                    set_cell(i_start, item.get("start_date", ""))
                    set_cell(i_end, item.get("end_date", ""))
                    set_cell(i_publisher, item.get("publisher", ""))
                    set_cell(i_targeting, item.get("targeting", ""))
                    set_cell(i_rate, item.get("rate", ""))
                    set_cell(i_metric, item.get("metric", "Impressions"))
                    set_cell(i_qty, item.get("quantity", ""))
                    set_cell(i_gross, item.get("gross_rate", ""))
                    if i_net is not None and item.get("net_rate"):
                        set_cell(i_net, item.get("net_rate", ""))

                return True

    return False


def fill_sales_contact_block(doc: "Document", sales_name: str, sales_email: str) -> None:
    if not (sales_name or sales_email):
        return

    fill_block_by_table_keyword(
        doc,
        keyword="Sales Contact",
        label_map=[
            ("Name:", sales_name),
            ("Sales Email:", sales_email),
            ("Email:", sales_email),
        ],
    )


def fill_special_instructions(doc: "Document", special_instructions: str, pt_size: int = 9, font_name: str = "Calibri") -> bool:
    if not special_instructions or not special_instructions.strip():
        return False

    si = special_instructions.strip()

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.lower().startswith("special instructions"):
            if ":" in t:
                p.text = "Special Instructions: " + si
            else:
                p.text = "Special Instructions"
                p2 = p._parent.add_paragraph(si)
                set_paragraph_font(p2, pt_size, font_name)
            set_paragraph_font(p, pt_size, font_name)
            return True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_txt = (cell.text or "").strip().lower()
                if "special instructions" in cell_txt:
                    if len(row.cells) >= 2 and cell is row.cells[0]:
                        row.cells[1].text = si
                        for p in row.cells[1].paragraphs:
                            set_paragraph_font(p, pt_size, font_name)
                        return True
                    else:
                        cell.paragraphs[-1].add_run("\n" + si)
                        for p in cell.paragraphs:
                            set_paragraph_font(p, pt_size, font_name)
                        return True
    return False


def generate_io_docx_bytes(
    template_path: Path,
    advertiser_name: str,
    advertiser_contact: str,
    agency_name: str,
    agency_contact: str,
    campaign_name: str,
    customer_ref: str,
    campaign_date: str,
    billing_customer_name: str,
    billing_contact_name: str,
    billing_address: str,
    billing_phone: str,
    billing_email: str,
    sales_contact_name: str,
    sales_contact_email: str,
    special_instructions: str,
    line_items: List[Dict],
    total_budget: float,
    total_impressions: int,
) -> bytes:
    doc = Document(str(template_path))

    set_document_styles(doc, pt_size=9, font_name="Calibri")

    fill_block_by_table_keyword(
        doc,
        keyword="Customer Information",
        label_map=[
            ("Advertiser Name:", advertiser_name),
            ("Advertiser Contact:", advertiser_contact),
            ("Agency Name:", agency_name),
            ("Agency Contact:", agency_contact),
        ],
    )

    fill_block_by_table_keyword(
        doc,
        keyword="Campaign Information",
        label_map=[
            ("Campaign Name:", campaign_name),
            ("Customer Reference Number:", customer_ref),
            ("Customer Job Number:", ""),
            ("Date:", campaign_date),
        ],
    )

    fill_block_by_table_keyword(
        doc,
        keyword="Billing Contact",
        label_map=[
            ("Customer Name:", billing_customer_name),
            ("Contact Name:", billing_contact_name),
            ("Address:", billing_address),
            ("Phone:", billing_phone),
            ("Billing Email:", billing_email),
            ("Email:", billing_email),
        ],
    )

    fill_sales_contact_block(doc, sales_contact_name, sales_contact_email)
    fill_special_instructions(doc, special_instructions, pt_size=9, font_name="Calibri")

    fill_media_buy_total_cell(doc, total_budget, pt_size=9, font_name="Calibri")
    fill_media_buy_rows(doc, line_items, pt_size=9, font_name="Calibri")
    fill_top_line_numbers(doc, total_budget, total_impressions, note="CPM: Mixed", pt_size=9, font_name="Calibri")

    set_document_styles(doc, pt_size=9, font_name="Calibri")

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


# ------------------------------------------------------------
# Budget extras helpers (Tab 2 -> Tab 3)
# ------------------------------------------------------------
def build_extras_calc(extra_df: pd.DataFrame) -> pd.DataFrame:
    if extra_df is None or extra_df.empty:
        return pd.DataFrame(columns=["Product", "Locations", "Rate (CPM)", "Quantity (Impressions)", "Gross Rate"])

    rows = []
    for _, r in extra_df.fillna(0).iterrows():
        prod = str(r.get("Product", "DOOH")).strip() or "DOOH"
        locs = int(r.get("Locations", 0) or 0)
        bud = float(r.get("Budget (R)", 0.0) or 0.0)

        cpm = dooh_cpm_from_count(locs) if prod == "DOOH" else mobile_cpm_from_count(locs)
        imps = impressions_from_budget_and_cpm(bud, cpm)

        rows.append({
            "Product": prod,
            "Locations": locs,
            "Rate (CPM)": f"R {cpm}",
            "Quantity (Impressions)": int(round(imps, 0)),
            "Gross Rate": float(bud),
        })

    return pd.DataFrame(rows)


def extras_fingerprint(extra_df: pd.DataFrame) -> str:
    try:
        payload = extra_df.fillna("").to_dict(orient="records") if isinstance(extra_df, pd.DataFrame) else []
        s = json.dumps(payload, sort_keys=True)
    except Exception:
        s = ""
    return hashlib.md5(s.encode("utf-8")).hexdigest()


def build_default_io_rows(
    campaign_type: str,
    dooh_budget: float,
    mobile_budget: float,
    dooh_count: int,
    mobile_count: int,
    start_date: str,
    end_date: str,
    extra_calc: Optional[pd.DataFrame],
) -> List[Dict]:
    rows = []

    base_idx = 0
    if campaign_type in ["DOOH", "DOOH + Mobile"]:
        dooh_cpm = dooh_cpm_from_count(dooh_count)
        dooh_imps = int(round(impressions_from_budget_and_cpm(dooh_budget, dooh_cpm), 0))
        rows.append({
            "_row_id": f"base_{base_idx}",
            "Product": "DOOH",
            "Description": "",
            "Start Date": start_date.strip(),
            "End Date": end_date.strip(),
            "Publisher": "",
            "Targeting": "",
            "Rate": f"R {dooh_cpm}",
            "Metric": "Impressions",
            "Quantity": fmt_int(dooh_imps),
            "Gross Rate": fmt_currency_rands(dooh_budget),
            "Net Rate (optional)": "",
            "_row_source": "base",
        })
        base_idx += 1

    if campaign_type in ["Mobile", "DOOH + Mobile"]:
        mob_cpm = mobile_cpm_from_count(mobile_count)
        mob_imps = int(round(impressions_from_budget_and_cpm(mobile_budget, mob_cpm), 0))
        rows.append({
            "_row_id": f"base_{base_idx}",
            "Product": "Mobile",
            "Description": "",
            "Start Date": start_date.strip(),
            "End Date": end_date.strip(),
            "Publisher": "",
            "Targeting": "",
            "Rate": f"R {mob_cpm}",
            "Metric": "Impressions",
            "Quantity": fmt_int(mob_imps),
            "Gross Rate": fmt_currency_rands(mobile_budget),
            "Net Rate (optional)": "",
            "_row_source": "base",
        })
        base_idx += 1

    if isinstance(extra_calc, pd.DataFrame) and not extra_calc.empty:
        for j, (_, r) in enumerate(extra_calc.iterrows()):
            prod = str(r.get("Product", "")).strip() or "DOOH"
            rate = str(r.get("Rate (CPM)", "")).strip()
            qty = fmt_int(r.get("Quantity (Impressions)", 0))
            gross = fmt_currency_rands(float(r.get("Gross Rate", 0.0)))
            rows.append({
                "_row_id": f"extra_{j}",
                "Product": prod,
                "Description": "",
                "Start Date": start_date.strip(),
                "End Date": end_date.strip(),
                "Publisher": "",
                "Targeting": "",
                "Rate": rate,
                "Metric": "Impressions",
                "Quantity": qty,
                "Gross Rate": gross,
                "Net Rate (optional)": "",
                "_row_source": "extra",
            })

    return rows


# ------------------------------------------------------------
# Background + header
# ------------------------------------------------------------
bg_path = find_asset("background")
if bg_path:
    inject_background(bg_path, white_overlay_opacity=0.45)

inject_app_css()

logo_path = find_asset("logo")
render_header(logo_path)


# ------------------------------------------------------------
# Tabs
# ------------------------------------------------------------
tab1, tab2, tab3 = st.tabs(["Proximity", "Budget & Selection", "Insertion Order (DOCX)"])


# ============================================================
# TAB 1: Proximity
# ============================================================
with tab1:
    st.subheader("Analysis type")
    mode = st.radio(
        "Select what you want to calculate",
        ["Store to DOOH", "Store to Store"],
        horizontal=True,
        key="mode_radio"
    )

    st.subheader("CSV template (download)")
    st.download_button(
        "Download Stores template (CSV)",
        data=csv_bytes_from_df(stores_template_df()),
        file_name="stores_template.csv",
        mime="text/csv",
        key="dl_stores_template_proximity_onlyone",
    )

    st.subheader("Upload files")

    dooh_df = None
    stores_file = None
    stores_a_file = None
    stores_b_file = None

    # Store-to-store Province filter
    store_store_selected_provinces = None

    if mode == "Store to DOOH":
        stores_file = st.file_uploader("Upload Stores (CSV)", type="csv", key="stores_dooh")

        dooh_df = load_dooh_master()
        if dooh_df is None:
            st.info("DOOH master not found. Temporary option: upload a DOOH CSV for this run.")
            dooh_file = st.file_uploader("Upload DOOH (CSV) [temporary]", type="csv", key="dooh_temp")
            if dooh_file:
                dooh_df = read_csv_uploaded(dooh_file)
                prov_col_tmp = pick_province_col(dooh_df)
                if prov_col_tmp:
                    dooh_df[prov_col_tmp] = dooh_df[prov_col_tmp].apply(normalize_province_name)
        else:
            prov_col = pick_province_col(dooh_df)
            if prov_col:
                provinces_live = sorted([p for p in dooh_df[prov_col].dropna().unique().tolist() if str(p).strip() != ""])
                st.caption(f"Loaded DOOH master: {DOOH_MASTER_PATH.name} | Provinces: {len(provinces_live)}")
    else:
        colA, colB = st.columns(2)
        with colA:
            stores_a_file = st.file_uploader("Upload Store List A (CSV)", type="csv", key="stores_a")
        with colB:
            stores_b_file = st.file_uploader("Upload Store List B (CSV)", type="csv", key="stores_b")

        # FIX (Cloud): preview from bytes (no seek dependency) so the Province selector reliably shows
        if stores_a_file is not None:
            try:
                preview_a = read_csv_uploaded(stores_a_file, nrows=250)
                prov_a = pick_province_col(preview_a)
            except Exception:
                prov_a = None
                preview_a = None

            if prov_a:
                try:
                    prov_vals = (
                        preview_a[prov_a]
                        .apply(normalize_province_name)
                        .astype(str)
                        .str.strip()
                    )
                    all_provs = sorted([p for p in prov_vals.unique().tolist() if p and p.lower() != "nan"])
                    if all_provs:
                        st.subheader("Store filters")
                        SELECT_ALL = "Select all provinces"
                        province_options = [SELECT_ALL] + all_provs
                        chosen = st.multiselect(
                            "Provinces (Store List A)",
                            options=province_options,
                            default=[SELECT_ALL],
                            key="store_store_province_filter_ui",
                        )
                        store_store_selected_provinces = all_provs if SELECT_ALL in chosen else chosen
                except Exception:
                    pass

    selected_provinces = None
    if mode == "Store to DOOH" and isinstance(dooh_df, pd.DataFrame) and not dooh_df.empty:
        prov_col = pick_province_col(dooh_df)
        if prov_col:
            all_provinces = sorted([p for p in dooh_df[prov_col].dropna().unique().tolist() if str(p).strip() != ""])

            st.subheader("DOOH filters")

            SELECT_ALL = "Select all provinces"
            province_options = [SELECT_ALL] + all_provinces

            chosen = st.multiselect(
                "Provinces",
                options=province_options,
                default=[SELECT_ALL],
                key="dooh_province_filter_ui",
            )

            if SELECT_ALL in chosen:
                selected_provinces = all_provinces
            else:
                selected_provinces = chosen

            if selected_provinces is not None and len(selected_provinces) > 0:
                dooh_df = dooh_df[dooh_df[prov_col].isin(selected_provinces)].copy()
        else:
            st.warning("DOOH master does not have a Province column, so the Province filter cannot be shown.")

    st.subheader("Settings")
    radius = st.slider("Radius (km)", 1, 50, 10, key="radius_slider")

    show_summary = st.checkbox("Show summary table", value=True, key="show_summary")
    include_pairwise = st.checkbox("Generate pairwise rows (within radius)", value=True, key="include_pairwise")

    run = st.button("Run proximity", key="run_prox")

    if run:
        if mode == "Store to DOOH":
            if stores_file is None:
                st.error("Please upload the Stores CSV first.")
                st.stop()
            if dooh_df is None:
                st.error("DOOH data not available. Add data/dooh_master.csv or upload a temporary DOOH CSV above.")
                st.stop()

            stores = read_csv_uploaded(stores_file)
            targets = dooh_df.copy()
            target_label = "DOOH"

        else:
            if (stores_a_file is None) or (stores_b_file is None):
                st.error("Please upload BOTH Store List A and Store List B.")
                st.stop()

            stores = read_csv_uploaded(stores_a_file)
            targets = read_csv_uploaded(stores_b_file)
            target_label = "Store"

        store_lat, store_lon = pick_lat_lon(stores)
        store_name_col = pick_name_col(stores)
        store_prov_col = pick_province_col(stores)

        if store_lat is None or store_lon is None:
            st.error("Could not detect latitude/longitude in Stores. Ensure columns include 'lat' and 'lon' (or 'lng').")
            st.stop()

        tgt_lat, tgt_lon = pick_lat_lon(targets)
        tgt_name_col = pick_name_col(targets)
        tgt_prov_col = pick_province_col(targets)

        if tgt_lat is None or tgt_lon is None:
            st.error(f"Could not detect latitude/longitude in {target_label} dataset.")
            st.stop()

        dooh_site_id_col = None
        dooh_network_col = None
        dooh_prov_col = None
        if mode == "Store to DOOH":
            dooh_site_id_col = pick_first_existing(
                targets,
                ["Site ID/Number", "Site ID", "SiteID", "Site Number", "Site_Number", "SiteID/Number", "SiteId", "Site"]
            )
            dooh_network_col = pick_first_existing(
                targets,
                ["Network:", "Network", "Network Type", "Network_Type", "NetworkType"]
            )
            dooh_prov_col = pick_province_col(targets)

        stores[store_lat] = pd.to_numeric(stores[store_lat], errors="coerce")
        stores[store_lon] = pd.to_numeric(stores[store_lon], errors="coerce")
        targets[tgt_lat] = pd.to_numeric(targets[tgt_lat], errors="coerce")
        targets[tgt_lon] = pd.to_numeric(targets[tgt_lon], errors="coerce")

        stores = stores.dropna(subset=[store_lat, store_lon]).reset_index(drop=True)
        targets = targets.dropna(subset=[tgt_lat, tgt_lon]).reset_index(drop=True)

        if store_prov_col:
            stores[store_prov_col] = stores[store_prov_col].apply(normalize_province_name)
        if tgt_prov_col:
            targets[tgt_prov_col] = targets[tgt_prov_col].apply(normalize_province_name)
        if dooh_prov_col:
            targets[dooh_prov_col] = targets[dooh_prov_col].apply(normalize_province_name)

        # Apply Store-to-Store Province filter
        if mode == "Store to Store" and store_prov_col and store_store_selected_provinces:
            stores = stores[stores[store_prov_col].isin(store_store_selected_provinces)].copy().reset_index(drop=True)

        if stores.empty or targets.empty:
            st.error("After cleaning coordinates, one of the datasets has no valid lat/lon rows.")
            st.stop()

        tgt_lats = targets[tgt_lat].to_numpy()
        tgt_lons = targets[tgt_lon].to_numpy()

        summary_rows = []
        pairwise_rows = []

        for i, s in stores.iterrows():
            s_name = str(s[store_name_col]) if (store_name_col and store_name_col in stores.columns and pd.notna(s[store_name_col])) else f"Location {i}"
            s_prov = ""
            if store_prov_col and store_prov_col in stores.columns and pd.notna(s.get(store_prov_col)):
                s_prov = str(s.get(store_prov_col))

            dists = haversine_vec(s[store_lat], s[store_lon], tgt_lats, tgt_lons)
            within_mask = dists <= radius
            idxs = np.where(within_mask)[0]

            nearest_idx = int(np.nanargmin(dists)) if len(dists) else None
            nearest_dist = float(dists[nearest_idx]) if nearest_idx is not None else np.nan

            if show_summary:
                row = {
                    "Location": s_name,
                    "Province": s_prov if s_prov else "",
                    f"{target_label} sites within {radius}km": int(len(idxs)),
                    f"Nearest {target_label} distance (km)": round(nearest_dist, 3) if np.isfinite(nearest_dist) else None,
                }

                if mode == "Store to DOOH":
                    row["Nearest DOOH Site ID/Number"] = (
                        str(targets.iloc[nearest_idx][dooh_site_id_col])
                        if (nearest_idx is not None and dooh_site_id_col and pd.notna(targets.iloc[nearest_idx][dooh_site_id_col]))
                        else ""
                    )
                    row["Nearest DOOH Network"] = (
                        str(targets.iloc[nearest_idx][dooh_network_col])
                        if (nearest_idx is not None and dooh_network_col and pd.notna(targets.iloc[nearest_idx][dooh_network_col]))
                        else ""
                    )
                    if dooh_prov_col:
                        row["Nearest DOOH Province"] = (
                            str(targets.iloc[nearest_idx][dooh_prov_col])
                            if (nearest_idx is not None and pd.notna(targets.iloc[nearest_idx].get(dooh_prov_col)))
                            else ""
                        )
                else:
                    row["Nearest Store"] = (
                        str(targets.iloc[nearest_idx][tgt_name_col])
                        if (nearest_idx is not None and tgt_name_col and tgt_name_col in targets.columns and pd.notna(targets.iloc[nearest_idx][tgt_name_col]))
                        else ""
                    )
                    if tgt_prov_col:
                        row["Nearest Store Province"] = (
                            str(targets.iloc[nearest_idx][tgt_prov_col])
                            if (nearest_idx is not None and pd.notna(targets.iloc[nearest_idx].get(tgt_prov_col)))
                            else ""
                        )

                summary_rows.append(row)

            if include_pairwise:
                for j in idxs:
                    pr = {
                        "Location": s_name,
                        "Province": s_prov if s_prov else "",
                        "Distance (km)": round(float(dists[j]), 3),
                    }

                    if mode == "Store to DOOH":
                        pr["Site ID/Number"] = (
                            str(targets.iloc[j][dooh_site_id_col])
                            if (dooh_site_id_col and pd.notna(targets.iloc[j][dooh_site_id_col]))
                            else ""
                        )
                        pr["Network"] = (
                            str(targets.iloc[j][dooh_network_col])
                            if (dooh_network_col and pd.notna(targets.iloc[j][dooh_network_col]))
                            else ""
                        )
                        if dooh_prov_col:
                            pr["DOOH Province"] = (
                                str(targets.iloc[j][dooh_prov_col])
                                if pd.notna(targets.iloc[j].get(dooh_prov_col))
                                else ""
                            )
                    else:
                        pr["Store"] = (
                            str(targets.iloc[j][tgt_name_col])
                            if (tgt_name_col and tgt_name_col in targets.columns and pd.notna(targets.iloc[j][tgt_name_col]))
                            else f"Store {j}"
                        )
                        if tgt_prov_col:
                            pr["Store Province"] = (
                                str(targets.iloc[j][tgt_prov_col])
                                if pd.notna(targets.iloc[j].get(tgt_prov_col))
                                else ""
                            )

                    pairwise_rows.append(pr)

        summary_df = pd.DataFrame(summary_rows) if show_summary else None
        pairwise_df = pd.DataFrame(pairwise_rows) if include_pairwise else None

        st.session_state["prox_results"] = {
            "mode": mode,
            "radius": radius,
            "summary_df": summary_df,
            "pairwise_df": pairwise_df,
        }

        if isinstance(summary_df, pd.DataFrame) and not summary_df.empty:
            st.subheader("Planning insights")
            stats = compute_insight_stats(summary_df, pairwise_df, mode, radius)
            if stats:
                if mode == "Store to DOOH":
                    title, body, bullets = build_vicinity_recommendation_dooh(stats)
                else:
                    title, body, bullets = build_vicinity_recommendation_store_to_store(stats)

                render_recommendation_box(title, body, bullets)

                with st.expander("Opportunity windows (locations with zero nearby sites)", expanded=False):
                    st.write(stats["opportunity_locations_preview"])

                if mode == "Store to Store" and stats.get("isolated_detail"):
                    with st.expander("Isolated stores: nearest store and distance (even if outside radius)", expanded=False):
                        st.dataframe(pd.DataFrame(stats["isolated_detail"]), use_container_width=True, hide_index=True)

                if mode == "Store to DOOH" and isinstance(stats.get("top_networks"), pd.DataFrame):
                    with st.expander("Top DOOH networks near your locations (unique sites)", expanded=False):
                        st.dataframe(stats["top_networks"], use_container_width=True)

        if show_summary:
            st.subheader("Results (summary)")
            st.dataframe(summary_df, use_container_width=True, hide_index=True)
            st.download_button(
                "Download summary (CSV)",
                summary_df.to_csv(index=False).encode("utf-8"),
                file_name="proximity_summary.csv",
                mime="text/csv",
                key="dl_prox_summary_csv",
            )

        if include_pairwise:
            st.subheader("Results (pairwise within radius)")
            st.dataframe(pairwise_df, use_container_width=True, hide_index=True)
            st.download_button(
                "Download pairwise (CSV)",
                pairwise_df.to_csv(index=False).encode("utf-8"),
                file_name="location_proximity_results.csv",
                mime="text/csv",
                key="dl_prox_pairwise_csv",
            )

        if (not show_summary) and (not include_pairwise):
            st.info("Select at least one output option (summary and/or pairwise).")


# ============================================================
# TAB 2: Budget & Selection
# ============================================================
with tab2:
    st.subheader("Budget, selection, CPM & impressions")

    campaign_type = st.radio(
        "Campaign type",
        ["DOOH", "Mobile", "DOOH + Mobile"],
        horizontal=True,
        key="campaign_type"
    )

    dooh_selected_count = int(st.session_state.get("selected_sites_count_dooh", 0))
    mobile_selected_count = int(st.session_state.get("selected_sites_count_mobile", 0))

    dooh_budget = float(st.session_state.get("budget_dooh", 200000.0))
    mobile_budget = float(st.session_state.get("budget_mobile", 150000.0))

    base_total_budget = 0.0
    base_total_imps = 0.0

    dooh_uploaded = False
    mobile_uploaded = False

    if campaign_type == "DOOH":
        st.caption("DOOH selection upload must include: 'Site ID/Number' and 'Selected' (1/0).")

        st.download_button(
            "Download DOOH selection template (CSV)",
            data=csv_bytes_from_df(dooh_selection_template_df()),
            file_name="dooh_selection_template.csv",
            mime="text/csv",
            key="dl_dooh_selection_template_tab2",
        )

        selection_file_dooh = st.file_uploader("Upload DOOH selected sites (CSV)", type="csv", key="selection_upload_dooh")
        dooh_uploaded = selection_file_dooh is not None

        dooh_budget = st.number_input("DOOH budget (R)", min_value=0.0, value=float(dooh_budget), step=1000.0, key="budget_dooh_input")

        if selection_file_dooh:
            try:
                _df = load_selected_sites_from_upload(selection_file_dooh)
                dooh_selected_count = int(len(_df))
                st.success(f"Loaded {dooh_selected_count} selected DOOH sites.")
            except Exception as e:
                st.error(str(e))
                dooh_selected_count = 0

        dooh_cpm = dooh_cpm_from_count(dooh_selected_count)
        dooh_imps = impressions_from_budget_and_cpm(dooh_budget, dooh_cpm)

        if dooh_uploaded and dooh_selected_count < 50:
            msg = savings_message("DOOH", dooh_selected_count, float(dooh_budget))
            if msg:
                st.info(msg)

        st.session_state["selected_sites_count_dooh"] = dooh_selected_count
        st.session_state["budget_dooh"] = float(dooh_budget)

        base_total_budget = float(dooh_budget)
        base_total_imps = float(dooh_imps)

        st.subheader("Summary")
        kpi_grid([
            {"label": "Selected DOOH sites", "value": str(dooh_selected_count)},
            {"label": "DOOH CPM", "value": f"R {dooh_cpm}"},
            {"label": "Estimated DOOH impressions", "value": fmt_int(dooh_imps)},
            {"label": "DOOH budget", "value": fmt_currency_rands(dooh_budget)},
        ])

    elif campaign_type == "Mobile":
        st.caption("Mobile upload supports a location list with lat/lon (Province optional).")

        st.download_button(
            "Download Mobile/Locations template (CSV)",
            data=csv_bytes_from_df(mobile_locations_template_df()),
            file_name="mobile_locations_template.csv",
            mime="text/csv",
            key="dl_mobile_locations_template_tab2",
        )

        selection_file_mobile = st.file_uploader("Upload Mobile locations (CSV)", type="csv", key="selection_upload_mobile")
        mobile_uploaded = selection_file_mobile is not None

        mobile_budget = st.number_input("Mobile budget (R)", min_value=0.0, value=float(mobile_budget), step=1000.0, key="budget_mobile_input")

        if selection_file_mobile:
            try:
                mobile_selected_count = count_mobile_locations_any_format(selection_file_mobile)
                st.success(f"Loaded {mobile_selected_count} Mobile locations.")
            except Exception as e:
                st.error(str(e))
                mobile_selected_count = 0

        mob_cpm = mobile_cpm_from_count(mobile_selected_count)
        mob_imps = impressions_from_budget_and_cpm(mobile_budget, mob_cpm)

        if mobile_uploaded and mobile_selected_count < 50:
            msg = savings_message("Mobile", mobile_selected_count, float(mobile_budget))
            if msg:
                st.info(msg)

        st.session_state["selected_sites_count_mobile"] = mobile_selected_count
        st.session_state["budget_mobile"] = float(mobile_budget)

        base_total_budget = float(mobile_budget)
        base_total_imps = float(mob_imps)

        st.subheader("Summary")
        kpi_grid([
            {"label": "Mobile locations", "value": str(mobile_selected_count)},
            {"label": "Mobile CPM", "value": f"R {mob_cpm}"},
            {"label": "Estimated Mobile impressions", "value": fmt_int(mob_imps)},
            {"label": "Mobile budget", "value": fmt_currency_rands(mobile_budget)},
        ])

    else:
        colA, colB = st.columns(2)
        with colA:
            st.caption("DOOH selection upload: 'Site ID/Number' + 'Selected'")
            st.download_button(
                "Download DOOH selection template (CSV)",
                data=csv_bytes_from_df(dooh_selection_template_df()),
                file_name="dooh_selection_template.csv",
                mime="text/csv",
                key="dl_dooh_selection_template_tab2_combo",
            )
            selection_file_dooh = st.file_uploader("Upload DOOH selected sites (CSV)", type="csv", key="selection_upload_dooh_combo")
            dooh_uploaded = selection_file_dooh is not None

            dooh_budget = st.number_input("DOOH budget (R)", min_value=0.0, value=float(dooh_budget), step=1000.0, key="budget_dooh_input_combo")

        with colB:
            st.caption("Mobile upload: location list with lat/lon (Province optional).")
            st.download_button(
                "Download Mobile/Locations template (CSV)",
                data=csv_bytes_from_df(mobile_locations_template_df()),
                file_name="mobile_locations_template.csv",
                mime="text/csv",
                key="dl_mobile_locations_template_tab2_combo",
            )
            selection_file_mobile = st.file_uploader("Upload Mobile locations (CSV)", type="csv", key="selection_upload_mobile_combo")
            mobile_uploaded = selection_file_mobile is not None

            mobile_budget = st.number_input("Mobile budget (R)", min_value=0.0, value=float(mobile_budget), step=1000.0, key="budget_mobile_input_combo")

        if selection_file_dooh:
            try:
                _df = load_selected_sites_from_upload(selection_file_dooh)
                dooh_selected_count = int(len(_df))
                st.success(f"Loaded {dooh_selected_count} selected DOOH sites.")
            except Exception as e:
                st.error(f"DOOH: {e}")
                dooh_selected_count = 0

        if selection_file_mobile:
            try:
                mobile_selected_count = count_mobile_locations_any_format(selection_file_mobile)
                st.success(f"Loaded {mobile_selected_count} Mobile locations.")
            except Exception as e:
                st.error(f"Mobile: {e}")
                mobile_selected_count = 0

        dooh_cpm = dooh_cpm_from_count(dooh_selected_count)
        mob_cpm = mobile_cpm_from_count(mobile_selected_count)

        dooh_imps = impressions_from_budget_and_cpm(dooh_budget, dooh_cpm)
        mob_imps = impressions_from_budget_and_cpm(mobile_budget, mob_cpm)

        if dooh_uploaded and dooh_selected_count < 50:
            msg_d = savings_message("DOOH", dooh_selected_count, float(dooh_budget))
            if msg_d:
                st.info(msg_d)

        if mobile_uploaded and mobile_selected_count < 50:
            msg_m = savings_message("Mobile", mobile_selected_count, float(mobile_budget))
            if msg_m:
                st.info(msg_m)

        total_budget = float(dooh_budget) + float(mobile_budget)
        total_imps = float(dooh_imps + mob_imps)

        st.session_state["selected_sites_count_dooh"] = dooh_selected_count
        st.session_state["selected_sites_count_mobile"] = mobile_selected_count
        st.session_state["budget_dooh"] = float(dooh_budget)
        st.session_state["budget_mobile"] = float(mobile_budget)

        base_total_budget = float(total_budget)
        base_total_imps = float(total_imps)

        st.subheader("Summary")
        kpi_grid([
            {"label": "DOOH sites", "value": str(dooh_selected_count)},
            {"label": "DOOH CPM", "value": f"R {dooh_cpm}"},
            {"label": "DOOH impressions", "value": fmt_int(dooh_imps)},
            {"label": "DOOH budget", "value": fmt_currency_rands(dooh_budget)},
            {"label": "Mobile locations", "value": str(mobile_selected_count)},
            {"label": "Mobile CPM", "value": f"R {mob_cpm}"},
            {"label": "Mobile impressions", "value": fmt_int(mob_imps)},
            {"label": "Mobile budget", "value": fmt_currency_rands(mobile_budget)},
            {"label": "Total impressions", "value": fmt_int(total_imps)},
            {"label": "Media Buy Total", "value": fmt_currency_rands(total_budget)},
            {"label": "Campaign type", "value": "DOOH + Mobile"},
            {"label": "CPM", "value": "Mixed"},
        ])

    st.divider()
    st.subheader("Additional line items (optional)")
    st.caption("Add extra DOOH/Mobile line items on top of the campaign budgets above. These will be included in the IO line items on Tab 3.")

    if "extra_lineitems_df" not in st.session_state:
        st.session_state["extra_lineitems_df"] = pd.DataFrame([
            {"Product": "DOOH", "Locations": 0, "Budget (R)": 0.0},
        ])

    btnA, btnB, _sp = st.columns([1, 1, 3])
    with btnA:
        if st.button("Add extra row", key="add_extra_row"):
            df = st.session_state["extra_lineitems_df"].copy()
            df.loc[len(df)] = {"Product": "DOOH", "Locations": 0, "Budget (R)": 0.0}
            st.session_state["extra_lineitems_df"] = df
    with btnB:
        if st.button("Remove last extra row", key="remove_extra_row"):
            df = st.session_state["extra_lineitems_df"].copy()
            if len(df) > 1:
                st.session_state["extra_lineitems_df"] = df.iloc[:-1].reset_index(drop=True)

    extra_df = st.data_editor(
        st.session_state["extra_lineitems_df"],
        num_rows="fixed",
        use_container_width=True,
        key="extra_lineitems_editor",
        column_config={
            "Product": st.column_config.SelectboxColumn("Product", options=["DOOH", "Mobile"]),
            "Locations": st.column_config.NumberColumn("Number of locations/DOOH placements", min_value=0, step=1),
            "Budget (R)": st.column_config.NumberColumn("Budget (R)", min_value=0.0, step=1000.0),
        },
    )
    st.session_state["extra_lineitems_df"] = extra_df

    for _, r in extra_df.fillna(0).iterrows():
        prod = str(r.get("Product", "")).strip() or "DOOH"
        locs = int(r.get("Locations", 0) or 0)
        bud = float(r.get("Budget (R)", 0.0) or 0.0)
        if locs > 0 and bud > 0 and locs < 50:
            msg = savings_message(prod, locs, bud)
            if msg:
                st.warning(msg)

    extra_calc = build_extras_calc(extra_df)
    st.session_state["extra_lineitems_calc"] = extra_calc
    st.session_state["extra_lineitems_fingerprint"] = extras_fingerprint(extra_df)

    extra_total_budget = float(extra_calc["Gross Rate"].sum()) if not extra_calc.empty else 0.0
    extra_total_imps = float(extra_calc["Quantity (Impressions)"].sum()) if not extra_calc.empty else 0.0

    if not extra_calc.empty:
        st.subheader("Additional line items summary")
        view_df = extra_calc.copy()
        view_df["Gross Rate"] = view_df["Gross Rate"].apply(fmt_currency_rands)
        view_df["Quantity (Impressions)"] = view_df["Quantity (Impressions)"].apply(fmt_int)
        st.dataframe(view_df, use_container_width=True, hide_index=True)

        combined_budget = float(base_total_budget) + float(extra_total_budget)
        combined_imps = float(base_total_imps) + float(extra_total_imps)

        st.subheader("Combined totals (campaign + additional)")
        kpi_grid([
            {"label": "Base media buy total", "value": fmt_currency_rands(base_total_budget)},
            {"label": "Extra media buy total", "value": fmt_currency_rands(extra_total_budget)},
            {"label": "Combined media buy total", "value": fmt_currency_rands(combined_budget)},
            {"label": "Combined impressions", "value": fmt_int(combined_imps)},
        ])
    else:
        st.info("Add at least 1 additional row if you want extra line items included in the IO.")


# ============================================================
# TAB 3: Insertion Order (DOCX)
# ============================================================
with tab3:
    st.subheader("Insertion Order (DOCX)")

    default_template = find_default_io_template()

    if Document is None:
        st.warning("DOCX generation needs python-docx. Install it with: pip install python-docx")
        st.stop()

    if not default_template:
        st.error("No default IO template found in /data. Please place your template .docx in the data folder.")
        st.stop()

    st.caption(f"Using template from /data: {default_template.name}")

    st.divider()
    st.subheader("Client + campaign details")

    col1, col2 = st.columns(2)
    with col1:
        advertiser_name = st.text_input("Advertiser Name", value="")
        advertiser_contact = st.text_input("Advertiser Contact", value="")
        agency_name = st.text_input("Agency Name", value="")
        agency_contact = st.text_input("Agency Contact (email / contact)", value="")
    with col2:
        campaign_name = st.text_input("Campaign Name", value="")
        customer_ref = st.text_input("Customer Reference Number", value="")
        campaign_date = st.text_input("Date (DD/MM/YYYY)", value=date.today().strftime("%d/%m/%Y"))

    st.subheader("Billing details")
    b1, b2 = st.columns(2)
    with b1:
        billing_customer_name = st.text_input("Billing - Customer Name", value="")
        billing_contact_name = st.text_input("Billing - Contact Name", value="")
        billing_phone = st.text_input("Billing - Phone", value="")
    with b2:
        billing_email = st.text_input("Billing - Email", value="")
        billing_address = st.text_area("Billing - Address", value="", height=80)

    st.subheader("Sales contact")

    if "sales_contact_email_val" not in st.session_state:
        st.session_state["sales_contact_name_sel"] = SALES_NAMES[0] if SALES_NAMES else ""
        st.session_state["sales_contact_email_val"] = SALES_EMAIL_BY_NAME.get(st.session_state["sales_contact_name_sel"], "")

    def _sync_sales_email():
        nm = st.session_state.get("sales_contact_name_sel", "")
        st.session_state["sales_contact_email_val"] = SALES_EMAIL_BY_NAME.get(nm, "")

    s1, s2 = st.columns(2)
    with s1:
        sales_contact_name = st.selectbox(
            "Sales Contact",
            options=SALES_NAMES,
            key="sales_contact_name_sel",
            on_change=_sync_sales_email,
        )
    with s2:
        sales_contact_email = st.text_input(
            "Sales Contact - Email",
            value=st.session_state.get("sales_contact_email_val", ""),
            key="sales_contact_email_val",
            disabled=True,
        )

    # FIX: ensure the sentence renders on Streamlit Cloud (HTML callout)
    st.markdown(
        """
        <div class="io-callout">
          Please download the Insertion Order, sign it, and share it with your Sales Contact to confirm your campaign.
          Be sure to include the targeted location list and the selected DOOH placements, should you have it, when submitting.
          This ensures your campaign is ready for activation.
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.subheader("Special instructions")
    special_instructions = st.text_area("Special instructions (e.g., Landing Page Actions: Drive, Map)", value="", height=90)

    campaign_type = st.session_state.get("campaign_type", "DOOH")
    dooh_budget = float(st.session_state.get("budget_dooh", 0.0))
    mobile_budget = float(st.session_state.get("budget_mobile", 0.0))
    dooh_count = int(st.session_state.get("selected_sites_count_dooh", 0))
    mobile_count = int(st.session_state.get("selected_sites_count_mobile", 0))

    dooh_cpm = dooh_cpm_from_count(dooh_count) if campaign_type in ["DOOH", "DOOH + Mobile"] else 0
    mob_cpm = mobile_cpm_from_count(mobile_count) if campaign_type in ["Mobile", "DOOH + Mobile"] else 0

    extra_calc = st.session_state.get("extra_lineitems_calc")

    # ------------------------------------------------------------
    # Editable line items that persist + only rebuild when upstream changes
    # + FIX: force Start/End Date columns to string
    # ------------------------------------------------------------
    st.subheader("Line items (auto-generated from Budget & Selection)")

    upstream_sig = {
        "campaign_type": str(campaign_type),
        "dooh_budget": float(dooh_budget),
        "mobile_budget": float(mobile_budget),
        "dooh_count": int(dooh_count),
        "mobile_count": int(mobile_count),
        "extra_fp": str(st.session_state.get("extra_lineitems_fingerprint", "")),
    }
    upstream_fp = hashlib.md5(json.dumps(upstream_sig, sort_keys=True).encode("utf-8")).hexdigest()

    # Build defaults ONLY if not present or upstream changed
    if ("io_lineitems_df" not in st.session_state) or (st.session_state.get("io_upstream_fp") != upstream_fp):
        desired_rows = build_default_io_rows(
            campaign_type=campaign_type,
            dooh_budget=dooh_budget,
            mobile_budget=mobile_budget,
            dooh_count=dooh_count,
            mobile_count=mobile_count,
            start_date="",
            end_date="",
            extra_calc=extra_calc,
        )
        st.session_state["io_lineitems_df"] = pd.DataFrame(desired_rows)
        st.session_state["io_upstream_fp"] = upstream_fp

    # Keep hidden columns in the stored df, only edit visible df
    stored_df = st.session_state["io_lineitems_df"].copy()
    hidden_cols = [c for c in ["_row_source", "_row_id"] if c in stored_df.columns]
    visible_df = stored_df.drop(columns=hidden_cols, errors="ignore").copy()

    must_be_text_cols = ["Start Date", "End Date", "Description", "Publisher", "Targeting", "Net Rate (optional)"]
    for c in must_be_text_cols:
        if c not in visible_df.columns:
            visible_df[c] = ""
        visible_df[c] = visible_df[c].astype("string").fillna("")
        visible_df[c] = visible_df[c].replace(["<NA>", "nan", "NaN", "None"], "")

    edited_visible = st.data_editor(
        visible_df,
        use_container_width=True,
        hide_index=True,
        num_rows="fixed",
        key="io_lineitems_editor",
        column_config={
            "Start Date": st.column_config.TextColumn("Start Date", help="Type e.g. 01/02/2026"),
            "End Date": st.column_config.TextColumn("End Date", help="Type e.g. 28/02/2026"),
        },
        disabled=[
            "Product", "Rate", "Metric", "Quantity", "Gross Rate"
        ],
    )

    for c in must_be_text_cols:
        if c in edited_visible.columns:
            edited_visible[c] = edited_visible[c].astype("string").fillna("").replace(["<NA>", "nan", "NaN", "None"], "")

    for c in hidden_cols:
        edited_visible[c] = stored_df[c].values

    st.session_state["io_lineitems_df"] = edited_visible
    st.session_state["io_upstream_fp"] = upstream_fp
    lineitems_df = edited_visible

    # ------------------------------------------------------------
    # Totals
    # ------------------------------------------------------------
    def parse_currency(s):
        try:
            s = str(s).replace("R", "").replace(",", "").strip()
            return float(s)
        except Exception:
            return 0.0

    def parse_intish(s):
        try:
            s = str(s).replace(",", "").strip()
            return int(float(s))
        except Exception:
            return 0

    total_budget = float(lineitems_df.get("Gross Rate", pd.Series(dtype=str)).apply(parse_currency).sum()) if not lineitems_df.empty else 0.0
    total_impressions = int(lineitems_df.get("Quantity", pd.Series(dtype=str)).apply(parse_intish).sum()) if not lineitems_df.empty else 0

    st.subheader("Auto-calculated totals")
    kpi_grid([
        {"label": "Media Buy Total", "value": fmt_currency_rands(total_budget)},
        {"label": "Campaign type", "value": str(campaign_type)},
        {"label": "Total impressions", "value": fmt_int(total_impressions)},
        {"label": "DOOH CPM", "value": (f"R {dooh_cpm}" if dooh_cpm else "—")},
        {"label": "Mobile CPM", "value": (f"R {mob_cpm}" if mob_cpm else "—")},
    ])

    gen = st.button("Generate IO (pre-populated DOCX)", key="btn_generate_io")

    if gen:
        try:
            line_items = []
            for _, r in lineitems_df.fillna("").iterrows():
                line_items.append({
                    "product": str(r.get("Product", "")).strip(),
                    "description": str(r.get("Description", "")).strip(),
                    "start_date": str(r.get("Start Date", "")).strip(),
                    "end_date": str(r.get("End Date", "")).strip(),
                    "publisher": str(r.get("Publisher", "")).strip(),
                    "targeting": str(r.get("Targeting", "")).strip(),
                    "rate": str(r.get("Rate", "")).strip(),
                    "metric": str(r.get("Metric", "Impressions")).strip() or "Impressions",
                    "quantity": str(r.get("Quantity", "")).strip(),
                    "gross_rate": str(r.get("Gross Rate", "")).strip(),
                    "net_rate": str(r.get("Net Rate (optional)", "")).strip(),
                })

            doc_bytes = generate_io_docx_bytes(
                template_path=default_template,
                advertiser_name=advertiser_name.strip(),
                advertiser_contact=advertiser_contact.strip(),
                agency_name=agency_name.strip(),
                agency_contact=agency_contact.strip(),
                campaign_name=campaign_name.strip(),
                customer_ref=customer_ref.strip(),
                campaign_date=campaign_date.strip(),

                billing_customer_name=billing_customer_name.strip(),
                billing_contact_name=billing_contact_name.strip(),
                billing_address=billing_address.strip(),
                billing_phone=billing_phone.strip(),
                billing_email=billing_email.strip(),

                sales_contact_name=sales_contact_name.strip(),
                sales_contact_email=sales_contact_email.strip(),

                special_instructions=special_instructions.strip(),
                line_items=line_items,
                total_budget=float(total_budget),
                total_impressions=int(total_impressions),
            )

            st.success("IO generated.")
            st.download_button(
                "Download IO (DOCX)",
                data=doc_bytes,
                file_name="Insertion_Order_Prepopulated.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_io_docx",
            )

        except Exception as e:
            st.error(f"Could not generate IO: {e}")


